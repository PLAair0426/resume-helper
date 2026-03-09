"""文档解析服务 - 支持PDF/DOCX/TXT多格式"""
import difflib
import html
import re
import logging
import zipfile
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


class ParserService:
    """简历文档解析，提取原始文本"""

    SKILL_KEYS = (
        "data_analysis",
        "visualization",
        "programming",
        "ai_tools",
        "design",
        "media",
        "languages",
        "other",
    )

    @staticmethod
    def _is_llm_json_parse_error(exc: Exception) -> bool:
        """Detect model-output formatting failures that should trigger fallback parsing."""
        detail = str(exc).strip().lower()
        markers = (
            "failed to parse llm json output",
            "jsondecodeerror",
            "expecting value: line 1 column 1",
            "extra data:",
            "unterminated string",
            "invalid control character",
            "llm returned empty content",
            "did not return valid json",
        )
        return any(marker in detail for marker in markers)

    @staticmethod
    def _safe_text(value: Any) -> str:
        if value is None:
            return ""
        return re.sub(r"\s+", " ", str(value)).strip()

    @staticmethod
    def _normalize_match_text(value: str) -> str:
        normalized = value.lower()
        normalized = re.sub(r"\s+", "", normalized)
        normalized = re.sub(
            r"[，。；：:,.!！?？·•\-_/\\|()（）\[\]{}<>《》\"'`~@#$%^&*+=]+",
            "",
            normalized,
        )
        return normalized.strip()

    def _dedupe_lines(self, lines: list[str]) -> list[str]:
        seen: set[str] = set()
        deduped: list[str] = []
        for line in lines:
            text = self._safe_text(line)
            if not text:
                continue
            key = self._normalize_match_text(text) or text.lower()
            if key in seen:
                continue
            seen.add(key)
            deduped.append(text)
        return deduped

    def _merge_text_sources(self, *sources: str) -> str:
        merged_lines: list[str] = []
        for source in sources:
            if not isinstance(source, str) or not source.strip():
                continue
            merged_lines.extend(
                [line.strip() for line in re.split(r"\r?\n", source) if line.strip()]
            )
        return "\n".join(self._dedupe_lines(merged_lines))

    @staticmethod
    def _collect_table_lines(tables: list[Any]) -> list[str]:
        lines: list[str] = []
        for table in tables or []:
            for row in getattr(table, "rows", []):
                cells: list[str] = []
                for cell in getattr(row, "cells", []):
                    text = ParserService._safe_text(getattr(cell, "text", ""))
                    if text:
                        cells.append(text)
                if cells:
                    lines.append(" | ".join(cells))
        return lines

    @staticmethod
    def _extract_lines_from_word_xml(xml_content: str) -> list[str]:
        lines: list[str] = []
        paragraph_blocks = re.findall(r"<w:p[\s\S]*?</w:p>", xml_content)
        if not paragraph_blocks:
            paragraph_blocks = [xml_content]

        for block in paragraph_blocks:
            fragments = re.findall(r"<w:t[^>]*>(.*?)</w:t>", block)
            if not fragments:
                continue
            text = "".join(html.unescape(fragment) for fragment in fragments)
            text = re.sub(r"\s+", " ", text).strip()
            if text:
                lines.append(text)
        return lines

    @staticmethod
    def _is_empty_value(value: Any) -> bool:
        if value is None:
            return True
        if isinstance(value, str):
            return not value.strip()
        if isinstance(value, (list, tuple, set, dict)):
            return len(value) == 0
        return False

    async def parse(
        self,
        file_path: str,
        file_format: str,
        api_key: Optional[str] = None,
        provider: Optional[str] = None,
        model: Optional[str] = None,
        api_endpoint: Optional[str] = None,
    ) -> dict:
        """
        解析简历文件，返回结构化数据。

        Returns:
            {
                "profile": dict (符合resume_profile.json schema),
                "confidence": float,
                "raw_text": str,
                "missing_fields": list,
                "open_questions": list,
            }
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 1. 提取原始文本
        raw_text = await self._extract_text(path, file_format)
        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("提取的文本内容过少，请检查文件是否正确")

        # 2. LLM结构化抽取
        from backend.services.llm_service import llm_service
        profile_result = await self._structured_extraction(
            llm_service,
            raw_text,
            api_key=api_key,
            provider=provider,
            model=model,
            api_endpoint=api_endpoint,
        )

        return profile_result

    async def _extract_text(self, path: Path, fmt: str) -> str:
        """根据文件格式提取文本"""
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "doc": self._extract_docx,
            "txt": self._extract_txt,
            "md": self._extract_txt,
        }
        extractor = extractors.get(fmt)
        if not extractor:
            raise ValueError(f"不支持的格式: {fmt}")
        return await extractor(path)

    async def _extract_pdf(self, path: Path) -> str:
        """PDF文本提取（pdfplumber优先，PyMuPDF降级）。"""
        # 1) pdfplumber
        try:
            import pdfplumber

            texts: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or page.extract_text(layout=True)
                    if text and text.strip():
                        texts.append(text.strip())
            merged = self._merge_text_sources("\n\n".join(texts))
            if merged:
                return merged
        except Exception as exc:
            logger.warning("pdfplumber extraction failed: %s", exc)

        # 2) 降级：PyMuPDF
        try:
            import fitz

            doc = fitz.open(str(path))
            texts: list[str] = []
            for page in doc:
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    texts.append(page_text.strip())
            doc.close()

            merged = self._merge_text_sources("\n\n".join(texts))
            if merged:
                return merged
        except ImportError:
            logger.warning("PyMuPDF not available")
        except Exception as exc:
            logger.warning("PyMuPDF extraction failed: %s", exc)

        raise RuntimeError("PDF文本提取失败，请尝试上传DOCX格式")

    async def _extract_docx(self, path: Path) -> str:
        """DOCX文本提取：段落 + 表格 + 页眉页脚 + XML文本框。"""
        doc_text = ""
        xml_text = ""

        # 1) python-docx 标准提取
        try:
            from docx import Document

            doc = Document(str(path))
            lines: list[str] = []

            lines.extend(
                [self._safe_text(p.text) for p in doc.paragraphs if self._safe_text(p.text)]
            )
            lines.extend(self._collect_table_lines(doc.tables))

            # 页眉页脚可能包含电话、邮箱、求职意向
            for section in doc.sections:
                header = getattr(section, "header", None)
                footer = getattr(section, "footer", None)
                if header:
                    lines.extend(
                        [
                            self._safe_text(p.text)
                            for p in header.paragraphs
                            if self._safe_text(p.text)
                        ]
                    )
                    lines.extend(self._collect_table_lines(header.tables))
                if footer:
                    lines.extend(
                        [
                            self._safe_text(p.text)
                            for p in footer.paragraphs
                            if self._safe_text(p.text)
                        ]
                    )
                    lines.extend(self._collect_table_lines(footer.tables))

            doc_text = "\n".join(self._dedupe_lines(lines))
        except Exception as exc:
            logger.warning("python-docx extraction failed: %s", exc)

        # 2) XML补充提取（WPS文本框、特殊布局）
        try:
            xml_text = await self._extract_docx_xml(path)
        except Exception as exc:
            logger.warning("DOCX XML extraction failed: %s", exc)

        merged = self._merge_text_sources(doc_text, xml_text)
        if merged:
            return merged

        raise RuntimeError("DOCX文本提取失败，请检查文件内容是否可读取")

    async def _extract_docx_xml(self, path: Path) -> str:
        """从DOCX XML中提取文本（处理WPS文本框/页眉页脚）。"""
        with zipfile.ZipFile(str(path), "r") as z:
            names = set(z.namelist())
            xml_files = ["word/document.xml"] + sorted(
                [
                    name
                    for name in names
                    if (name.startswith("word/header") or name.startswith("word/footer"))
                    and name.endswith(".xml")
                ]
            )

            lines: list[str] = []
            for xml_file in xml_files:
                if xml_file not in names:
                    continue
                xml_content = z.read(xml_file).decode("utf-8", errors="ignore")
                lines.extend(self._extract_lines_from_word_xml(xml_content))

        deduped = self._dedupe_lines(lines)
        if not deduped:
            raise ValueError("DOCX XML中未找到文本内容")
        return "\n".join(deduped)

    async def _extract_txt(self, path: Path) -> str:
        """纯文本提取"""
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "utf-16"]
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("无法识别文件编码")

    def _ensure_profile_shape(self, profile: Any) -> dict:
        """确保profile结构稳定，避免前端映射时丢字段。"""
        if not isinstance(profile, dict):
            profile = {}

        basics = profile.get("basics")
        if not isinstance(basics, dict):
            basics = {}

        def _pick_text(*values: Any) -> str:
            for value in values:
                text = self._safe_text(value)
                if text:
                    return text
            return ""

        def _period_bound(item: dict, edge: str) -> str:
            periods = item.get("periods")
            if not isinstance(periods, list):
                return ""
            candidates = [entry for entry in periods if isinstance(entry, dict)]
            if not candidates:
                return ""
            target = candidates[0] if edge == "start" else candidates[-1]
            alt = "from" if edge == "start" else "to"
            return _pick_text(target.get(edge), target.get(alt), target.get("date"))

        def _normalize_responsibilities(item: dict) -> list[dict]:
            normalized: list[dict] = []
            responsibilities = item.get("responsibilities")
            if isinstance(responsibilities, list):
                for entry in responsibilities:
                    if isinstance(entry, dict):
                        description = _pick_text(
                            entry.get("description"),
                            entry.get("content"),
                            entry.get("text"),
                            entry.get("detail"),
                            entry.get("highlight"),
                        )
                        if not description:
                            continue
                        normalized.append(
                            {
                                "category": _pick_text(entry.get("category")),
                                "description": description,
                                "metrics": self._unique_str_list(
                                    entry.get("metrics", []) if isinstance(entry.get("metrics"), list) else []
                                ),
                                "keywords": self._unique_str_list(
                                    entry.get("keywords", []) if isinstance(entry.get("keywords"), list) else []
                                ),
                            }
                        )
                        continue

                    text = self._safe_text(entry)
                    if not text:
                        continue
                    normalized.append(
                        {
                            "category": "",
                            "description": text,
                            "metrics": [],
                            "keywords": [],
                        }
                    )

            if normalized:
                return self._merge_responsibilities(normalized, [])

            fallback_lines: list[str] = []
            for key in ("highlights", "achievements", "bullets"):
                value = item.get(key)
                if isinstance(value, list):
                    fallback_lines.extend([self._safe_text(v) for v in value if self._safe_text(v)])

            blob = _pick_text(item.get("details"), item.get("description"), item.get("summary"))
            if blob:
                fallback_lines.extend(
                    [self._safe_text(v) for v in re.split(r"[\r\n;；]+", blob) if self._safe_text(v)]
                )

            return [
                {
                    "category": "",
                    "description": line,
                    "metrics": [],
                    "keywords": [],
                }
                for line in self._unique_str_list(fallback_lines)
            ]

        education: list[dict] = []
        for item in profile.get("education", []):
            if not isinstance(item, dict):
                continue
            normalized = {
                "school": _pick_text(item.get("school"), item.get("university"), item.get("college")),
                "degree_type": _pick_text(item.get("degree_type"), item.get("degree")),
                "major": _pick_text(item.get("major"), item.get("field"), item.get("specialization")),
                "start_date": _pick_text(item.get("start_date"), item.get("start"), item.get("from")),
                "end_date": _pick_text(item.get("end_date"), item.get("end"), item.get("to")),
                "gpa": _pick_text(item.get("gpa")) or None,
                "courses": self._unique_str_list(item.get("courses", []) if isinstance(item.get("courses"), list) else []),
                "awards": self._unique_str_list(item.get("awards", []) if isinstance(item.get("awards"), list) else []),
            }
            has_content = any(
                [
                    normalized["school"],
                    normalized["degree_type"],
                    normalized["major"],
                    normalized["start_date"],
                    normalized["end_date"],
                    normalized["gpa"],
                    normalized["courses"],
                    normalized["awards"],
                ]
            )
            if has_content:
                education.append(normalized)

        work_experience: list[dict] = []
        for item in profile.get("work_experience", []):
            if not isinstance(item, dict):
                continue

            role = _pick_text(
                item.get("role"),
                item.get("title"),
                item.get("position"),
                item.get("job_title"),
            )
            company = _pick_text(
                item.get("company"),
                item.get("organization"),
                item.get("employer"),
            )
            start_date = _pick_text(
                item.get("start_date"),
                item.get("start"),
                item.get("from"),
                _period_bound(item, "start"),
            )
            end_date = _pick_text(
                item.get("end_date"),
                item.get("end"),
                item.get("to"),
                _period_bound(item, "end"),
            )
            responsibilities = _normalize_responsibilities(item)
            internship_raw = item.get("is_internship")
            is_internship = (
                internship_raw
                if isinstance(internship_raw, bool)
                else bool(re.search(r"(实习|intern)", role, re.IGNORECASE))
            )

            has_content = bool(
                company
                or role
                or start_date
                or end_date
                or responsibilities
            )
            if not has_content:
                continue

            work_experience.append(
                {
                    "company": company,
                    "role": role,
                    "start_date": start_date,
                    "end_date": end_date,
                    "is_internship": is_internship,
                    "responsibilities": responsibilities,
                }
            )

        research: list[dict] = []
        for item in profile.get("research", []):
            if not isinstance(item, dict):
                continue
            publications: list[dict] = []
            for pub in item.get("publications", []) if isinstance(item.get("publications"), list) else []:
                if not isinstance(pub, dict):
                    continue
                title = _pick_text(pub.get("title"), pub.get("name"))
                venue = _pick_text(pub.get("venue"), pub.get("journal"), pub.get("publisher"))
                status = _pick_text(pub.get("status"), pub.get("state"))
                if not (title or venue or status):
                    continue
                publications.append({"title": title, "venue": venue, "status": status})

            normalized = {
                "title": _pick_text(item.get("title"), item.get("name"), item.get("project")),
                "type": _pick_text(item.get("type"), item.get("category")),
                "start_date": _pick_text(item.get("start_date"), item.get("start"), item.get("from")),
                "end_date": _pick_text(item.get("end_date"), item.get("end"), item.get("to")),
                "description": _pick_text(item.get("description"), item.get("summary")),
                "contributions": _pick_text(item.get("contributions"), item.get("role")),
                "publications": publications,
            }

            has_content = any(
                [
                    normalized["title"],
                    normalized["type"],
                    normalized["start_date"],
                    normalized["end_date"],
                    normalized["description"],
                    normalized["contributions"],
                    normalized["publications"],
                ]
            )
            if has_content:
                research.append(normalized)

        skills = profile.get("skills")
        if not isinstance(skills, dict):
            skills = {}
        metadata = profile.get("metadata")
        if not isinstance(metadata, dict):
            metadata = {}

        for item in work_experience:
            responsibilities = item.get("responsibilities")
            if not isinstance(responsibilities, list):
                item["responsibilities"] = []
                continue
            item["responsibilities"] = [
                r for r in responsibilities if isinstance(r, dict) and self._safe_text(r.get("description"))
            ]

        for key in self.SKILL_KEYS:
            if key == "languages":
                value = skills.get(key)
                skills[key] = [v for v in value if isinstance(v, dict)] if isinstance(value, list) else []
            else:
                value = skills.get(key)
                skills[key] = [self._safe_text(v) for v in value if self._safe_text(v)] if isinstance(value, list) else []

        open_questions = metadata.get("open_questions")
        metadata["open_questions"] = (
            [self._safe_text(v) for v in open_questions if self._safe_text(v)]
            if isinstance(open_questions, list)
            else []
        )

        missing_fields = metadata.get("missing_fields")
        metadata["missing_fields"] = (
            [self._safe_text(v) for v in missing_fields if self._safe_text(v)]
            if isinstance(missing_fields, list)
            else []
        )

        return {
            "basics": {
                "name": _pick_text(basics.get("name"), basics.get("full_name")),
                "phone": _pick_text(basics.get("phone"), basics.get("mobile")) or None,
                "email": _pick_text(basics.get("email"), basics.get("mail")) or None,
                "degree": _pick_text(basics.get("degree"), basics.get("highest_degree")) or None,
                "graduation_year": basics.get("graduation_year"),
                "target_role": _pick_text(
                    basics.get("target_role"),
                    basics.get("target_position"),
                    basics.get("job_title"),
                )
                or None,
                "summary": _pick_text(basics.get("summary"), basics.get("profile")) or None,
            },
            "education": education,
            "work_experience": work_experience,
            "research": research,
            "skills": skills,
            "metadata": metadata,
        }

    def _unique_str_list(self, values: list[Any]) -> list[str]:
        output: list[str] = []
        seen: set[str] = set()
        for value in values:
            text = self._safe_text(value)
            if not text:
                continue
            key = self._normalize_match_text(text) or text
            if key in seen:
                continue
            seen.add(key)
            output.append(text)
        return output

    def _unique_language_list(self, values: list[Any]) -> list[dict]:
        output: list[dict] = []
        seen: set[str] = set()
        for value in values:
            if not isinstance(value, dict):
                continue
            language = self._safe_text(value.get("language"))
            level = self._safe_text(value.get("level"))
            if not language:
                continue
            key = f"{self._normalize_match_text(language)}::{self._normalize_match_text(level)}"
            if key in seen:
                continue
            seen.add(key)
            output.append({"language": language, "level": level})
        return output

    def _record_key(self, item: dict, fields: tuple[str, ...]) -> str:
        tokens = [self._normalize_match_text(self._safe_text(item.get(field))) for field in fields]
        tokens = [token for token in tokens if token]
        return "|".join(tokens)

    def _work_record_key(self, item: dict) -> str:
        """Build a conservative key to avoid over-merging distinct work entries."""
        company = self._normalize_match_text(self._safe_text(item.get("company")))
        role = self._normalize_match_text(self._safe_text(item.get("role")))
        start = self._normalize_match_text(self._safe_text(item.get("start_date")))
        end = self._normalize_match_text(self._safe_text(item.get("end_date")))

        if not company and not role:
            return ""

        if start or end:
            return "|".join(token for token in (company, role, start, end) if token)

        # Without dates, include leading responsibility fingerprints to avoid
        # collapsing multiple jobs into one item.
        fingerprints: list[str] = []
        responsibilities = (
            item.get("responsibilities")
            if isinstance(item.get("responsibilities"), list)
            else []
        )
        for resp in responsibilities:
            if not isinstance(resp, dict):
                continue
            desc = self._normalize_match_text(self._safe_text(resp.get("description")))
            if not desc:
                continue
            fingerprints.append(desc[:80])
            if len(fingerprints) >= 2:
                break

        base = "|".join(token for token in (company, role) if token)
        if fingerprints:
            return f"{base}|{'|'.join(fingerprints)}"
        return base

    def _merge_responsibilities(self, first: list[Any], second: list[Any]) -> list[dict]:
        merged: list[dict] = []
        seen: set[str] = set()

        for source in [first, second]:
            for item in source:
                if not isinstance(item, dict):
                    continue
                description = self._safe_text(item.get("description"))
                if not description:
                    continue
                key = self._normalize_match_text(description)
                if key in seen:
                    continue
                seen.add(key)
                merged.append(
                    {
                        "category": self._safe_text(item.get("category")),
                        "description": description,
                        "metrics": self._unique_str_list(item.get("metrics", []) if isinstance(item.get("metrics"), list) else []),
                        "keywords": self._unique_str_list(item.get("keywords", []) if isinstance(item.get("keywords"), list) else []),
                    }
                )
        return merged

    @staticmethod
    def _is_pure_date_like_line(text: str) -> bool:
        compact = re.sub(r"\s+", "", text)
        if not compact:
            return True
        return bool(
            re.fullmatch(
                r"[0-9一二三四五六七八九十年月日\.\-/~—–至到、,，（）()]+",
                compact,
            )
        )

    def _looks_like_company_header_line(self, text: str) -> bool:
        line = self._safe_text(text)
        if not line:
            return False
        if len(line) > 90:
            return False
        if re.search(
            r"(公司|集团|有限公司|股份|研究院|事务所|inc\.?|ltd\.?|corp\.?|llc)",
            line,
            re.IGNORECASE,
        ):
            return True
        return False

    def _responsibilities_from_lines(self, lines: list[str]) -> list[dict]:
        output: list[dict] = []
        for line in self._unique_str_list(lines):
            text = self._safe_text(line)
            if not text:
                continue
            output.append(
                {
                    "category": "",
                    "description": text,
                    "metrics": [],
                    "keywords": [],
                }
            )
        return output

    def _redistribute_work_responsibilities_from_raw(
        self,
        parsed: dict,
        raw_text: str,
    ) -> None:
        """Redistribute work responsibilities by company blocks from source text."""
        work_items = parsed.get("work_experience")
        if not isinstance(work_items, list) or len(work_items) < 2:
            return

        raw_lines = self._dedupe_lines(
            [self._safe_text(line) for line in raw_text.splitlines() if self._safe_text(line)]
        )
        if len(raw_lines) < 6:
            return

        companies: list[tuple[int, str, str]] = []
        for idx, item in enumerate(work_items):
            if not isinstance(item, dict):
                continue
            company = self._safe_text(item.get("company"))
            company_key = self._normalize_match_text(company)
            if len(company_key) < 3:
                continue
            companies.append((idx, company, company_key))

        if len(companies) < 2:
            return

        hits: list[tuple[int, int, str, str]] = []
        used_lines: set[int] = set()
        for exp_idx, company, company_key in companies:
            found_index = -1
            for line_idx, line in enumerate(raw_lines):
                if line_idx in used_lines:
                    continue
                normalized_line = self._normalize_match_text(line)
                if not normalized_line:
                    continue
                if company_key in normalized_line or normalized_line in company_key:
                    found_index = line_idx
                    break
            if found_index >= 0:
                used_lines.add(found_index)
                hits.append((exp_idx, found_index, company, company_key))

        if len(hits) < 2:
            return

        hits.sort(key=lambda item: item[1])
        block_lines_map: dict[int, list[str]] = {}
        for pos, (exp_idx, line_idx, _, company_key) in enumerate(hits):
            start = line_idx + 1
            end = hits[pos + 1][1] if pos + 1 < len(hits) else len(raw_lines)
            if start >= end:
                continue

            block_lines: list[str] = []
            for line in raw_lines[start:end]:
                text = self._safe_text(line)
                if not text:
                    continue
                normalized = self._normalize_match_text(text)
                if not normalized:
                    continue
                if company_key and (company_key in normalized or normalized in company_key):
                    continue
                if self._looks_like_company_header_line(text):
                    continue
                if self._is_pure_date_like_line(text):
                    continue
                if len(normalized) <= 3 and not any(ch.isdigit() for ch in normalized):
                    continue
                block_lines.append(text)

            cleaned_block = self._unique_str_list(block_lines)
            if cleaned_block:
                block_lines_map[exp_idx] = cleaned_block

        if not block_lines_map:
            return

        resp_counts: list[int] = []
        for item in work_items:
            if not isinstance(item, dict):
                resp_counts.append(0)
                continue
            responsibilities = (
                item.get("responsibilities")
                if isinstance(item.get("responsibilities"), list)
                else []
            )
            count = 0
            for resp in responsibilities:
                if isinstance(resp, dict) and self._safe_text(resp.get("description")):
                    count += 1
            resp_counts.append(count)

        total = sum(resp_counts)
        max_count = max(resp_counts) if resp_counts else 0
        empty_count = sum(1 for count in resp_counts if count == 0)
        avg = total / len(resp_counts) if resp_counts else 0.0
        is_imbalanced = empty_count >= 1 and max_count >= max(6, int(avg * 2))

        for exp_idx, lines in block_lines_map.items():
            if exp_idx < 0 or exp_idx >= len(work_items):
                continue
            item = work_items[exp_idx]
            if not isinstance(item, dict):
                continue

            block_responsibilities = self._responsibilities_from_lines(lines)
            if not block_responsibilities:
                continue

            existing = item.get("responsibilities")
            existing_list = existing if isinstance(existing, list) else []

            if is_imbalanced:
                item["responsibilities"] = self._merge_responsibilities(
                    block_responsibilities,
                    [],
                )
            elif not existing_list:
                item["responsibilities"] = self._merge_responsibilities(
                    block_responsibilities,
                    [],
                )
            else:
                item["responsibilities"] = self._merge_responsibilities(
                    existing_list,
                    block_responsibilities,
                )

    def _merge_profiles(self, primary: dict, secondary: dict) -> dict:
        """将首轮与补全轮抽取结果进行并集合并。"""
        merged = self._ensure_profile_shape(primary)
        supplement = self._ensure_profile_shape(secondary)

        # basics：优先保留primary，缺失时用supplement补齐
        for field in ("name", "phone", "email", "degree", "graduation_year", "target_role"):
            current = merged["basics"].get(field)
            candidate = supplement["basics"].get(field)
            if self._is_empty_value(current) and not self._is_empty_value(candidate):
                merged["basics"][field] = candidate

        summary_current = self._safe_text(merged["basics"].get("summary"))
        summary_candidate = self._safe_text(supplement["basics"].get("summary"))
        if len(summary_candidate) > len(summary_current):
            merged["basics"]["summary"] = summary_candidate

        # education 并集
        edu_map: dict[str, dict] = {}
        for item in merged["education"]:
            key = self._record_key(item, ("school", "major", "degree_type", "start_date", "end_date"))
            if key:
                edu_map[key] = item
        for item in supplement["education"]:
            key = self._record_key(item, ("school", "major", "degree_type", "start_date", "end_date"))
            if not key:
                item["responsibilities"] = self._merge_responsibilities(
                    item.get("responsibilities", []), []
                )
                merged["work_experience"].append(item)
                continue
            if key not in edu_map:
                merged["education"].append(item)
                edu_map[key] = item
                continue
            target = edu_map[key]
            for field in ("school", "degree_type", "major", "start_date", "end_date", "gpa"):
                if self._is_empty_value(target.get(field)) and not self._is_empty_value(item.get(field)):
                    target[field] = item.get(field)
            target["courses"] = self._unique_str_list(
                (target.get("courses") if isinstance(target.get("courses"), list) else [])
                + (item.get("courses") if isinstance(item.get("courses"), list) else [])
            )
            target["awards"] = self._unique_str_list(
                (target.get("awards") if isinstance(target.get("awards"), list) else [])
                + (item.get("awards") if isinstance(item.get("awards"), list) else [])
            )

        # work_experience 并集
        work_map: dict[str, dict] = {}
        for item in merged["work_experience"]:
            key = self._work_record_key(item)
            if key:
                work_map[key] = item
        for item in supplement["work_experience"]:
            key = self._work_record_key(item)
            if not key:
                item["responsibilities"] = self._merge_responsibilities(
                    item.get("responsibilities", []), []
                )
                merged["work_experience"].append(item)
                continue
            if key not in work_map:
                item["responsibilities"] = self._merge_responsibilities(
                    item.get("responsibilities", []), []
                )
                merged["work_experience"].append(item)
                work_map[key] = item
                continue

            target = work_map[key]
            for field in ("company", "role", "start_date", "end_date", "is_internship"):
                if self._is_empty_value(target.get(field)) and not self._is_empty_value(item.get(field)):
                    target[field] = item.get(field)

            target["responsibilities"] = self._merge_responsibilities(
                target.get("responsibilities", []),
                item.get("responsibilities", []),
            )

        # research 并集
        research_map: dict[str, dict] = {}
        for item in merged["research"]:
            key = self._record_key(item, ("title", "start_date", "end_date"))
            if key:
                research_map[key] = item
        for item in supplement["research"]:
            key = self._record_key(item, ("title", "start_date", "end_date"))
            if not key:
                continue
            if key not in research_map:
                merged["research"].append(item)
                research_map[key] = item
                continue
            target = research_map[key]
            for field in (
                "title",
                "type",
                "start_date",
                "end_date",
                "description",
                "contributions",
            ):
                if self._is_empty_value(target.get(field)) and not self._is_empty_value(item.get(field)):
                    target[field] = item.get(field)

            target["publications"] = [
                pub
                for pub in (target.get("publications") if isinstance(target.get("publications"), list) else [])
                if isinstance(pub, dict)
            ]
            pubs = item.get("publications") if isinstance(item.get("publications"), list) else []
            existing_keys = {
                self._record_key(pub, ("title", "venue", "status"))
                for pub in target["publications"]
                if isinstance(pub, dict)
            }
            for pub in pubs:
                if not isinstance(pub, dict):
                    continue
                pub_key = self._record_key(pub, ("title", "venue", "status"))
                if pub_key and pub_key not in existing_keys:
                    target["publications"].append(pub)
                    existing_keys.add(pub_key)

        # skills 并集
        all_skill_keys = set(merged["skills"].keys()) | set(supplement["skills"].keys()) | set(self.SKILL_KEYS)
        for key in all_skill_keys:
            current = merged["skills"].get(key, [])
            candidate = supplement["skills"].get(key, [])
            if key == "languages":
                merged["skills"][key] = self._unique_language_list(
                    (current if isinstance(current, list) else [])
                    + (candidate if isinstance(candidate, list) else [])
                )
            else:
                merged["skills"][key] = self._unique_str_list(
                    (current if isinstance(current, list) else [])
                    + (candidate if isinstance(candidate, list) else [])
                )

        # metadata 合并
        meta = merged.setdefault("metadata", {})
        meta2 = supplement.get("metadata", {})
        if not isinstance(meta2, dict):
            meta2 = {}

        meta["open_questions"] = self._unique_str_list(
            (meta.get("open_questions") if isinstance(meta.get("open_questions"), list) else [])
            + (meta2.get("open_questions") if isinstance(meta2.get("open_questions"), list) else [])
        )
        meta["missing_fields"] = self._unique_str_list(
            (meta.get("missing_fields") if isinstance(meta.get("missing_fields"), list) else [])
            + (meta2.get("missing_fields") if isinstance(meta2.get("missing_fields"), list) else [])
        )
        if self._is_empty_value(meta.get("source_format")) and not self._is_empty_value(meta2.get("source_format")):
            meta["source_format"] = meta2.get("source_format")
        if self._is_empty_value(meta.get("version")) and not self._is_empty_value(meta2.get("version")):
            meta["version"] = meta2.get("version")

        conf_current = meta.get("extraction_confidence")
        conf_candidate = meta2.get("extraction_confidence")
        try:
            conf_current = float(conf_current)
        except (TypeError, ValueError):
            conf_current = 0.75
        try:
            conf_candidate = float(conf_candidate)
        except (TypeError, ValueError):
            conf_candidate = conf_current
        meta["extraction_confidence"] = round(max(conf_current, conf_candidate), 3)

        return merged

    def _compute_missing_fields(self, parsed: dict) -> list[str]:
        missing: list[str] = []
        basics = parsed.get("basics", {})
        if self._is_empty_value(basics.get("name")):
            missing.append("basics.name")
        if self._is_empty_value(basics.get("phone")):
            missing.append("basics.phone")
        if self._is_empty_value(basics.get("email")):
            missing.append("basics.email")

        if not parsed.get("education"):
            missing.append("education")
        if not parsed.get("work_experience"):
            missing.append("work_experience")
        if not parsed.get("research"):
            missing.append("research")

        skills = parsed.get("skills", {})
        has_skill = False
        if isinstance(skills, dict):
            for value in skills.values():
                if isinstance(value, list) and value:
                    has_skill = True
                    break
        if not has_skill:
            missing.append("skills")

        return self._unique_str_list(missing)

    def _has_any_skill(self, parsed: dict) -> bool:
        skills = parsed.get("skills", {})
        if not isinstance(skills, dict):
            return False
        for value in skills.values():
            if isinstance(value, list) and value:
                return True
        return False

    def _is_fallback_parse(self, metadata: dict) -> bool:
        open_questions = metadata.get("open_questions")
        if not isinstance(open_questions, list):
            return False
        for item in open_questions:
            text = self._safe_text(item).lower()
            if "fallback" in text or "llm reason" in text:
                return True
        return False

    def _recalculate_confidence(
        self,
        parsed: dict,
        raw_text: str,
        missing_fields: list[str],
        unmapped_lines: list[str],
    ) -> float:
        metadata = parsed.get("metadata")
        if not isinstance(metadata, dict):
            metadata = {}

        base_confidence = metadata.get("extraction_confidence")
        try:
            base_value = float(base_confidence)
        except (TypeError, ValueError):
            base_value = 0.82

        basics = parsed.get("basics", {}) if isinstance(parsed.get("basics"), dict) else {}
        has_name = bool(self._safe_text(basics.get("name")))
        has_contact = bool(
            self._safe_text(basics.get("phone")) or self._safe_text(basics.get("email"))
        )
        has_education = bool(parsed.get("education"))
        has_work = bool(parsed.get("work_experience"))
        has_project = bool(parsed.get("research"))
        has_skill = self._has_any_skill(parsed)

        checks = [has_name, has_contact, has_education, has_work, has_project, has_skill]
        completeness = sum(1 for item in checks if item) / len(checks)

        # Full structured hit should be shown as 100%.
        if completeness >= 1.0 and not self._is_fallback_parse(metadata):
            return 1.0

        raw_line_count = len(
            [line for line in raw_text.splitlines() if self._safe_text(line)]
        )
        unmapped_ratio = len(unmapped_lines) / max(raw_line_count, 1)
        unmapped_penalty = min(0.05, unmapped_ratio * 0.05)
        missing_penalty = min(0.12, len(missing_fields) * 0.02)

        score = max(base_value, completeness)
        score = score - unmapped_penalty - missing_penalty

        if self._is_fallback_parse(metadata):
            score = min(score, 0.35)

        return round(max(0.0, min(score, 1.0)), 3)

    def _collect_profile_fragments(self, parsed: dict) -> list[str]:
        fragments: list[str] = []
        basics = parsed.get("basics", {})
        if isinstance(basics, dict):
            for field in (
                "name",
                "phone",
                "email",
                "degree",
                "graduation_year",
                "target_role",
                "summary",
            ):
                value = self._safe_text(basics.get(field))
                if value:
                    fragments.append(value)

        for edu in parsed.get("education", []):
            if not isinstance(edu, dict):
                continue
            for field in ("school", "degree_type", "major", "start_date", "end_date", "gpa"):
                value = self._safe_text(edu.get(field))
                if value:
                    fragments.append(value)
            fragments.extend(
                self._unique_str_list(edu.get("courses", []) if isinstance(edu.get("courses"), list) else [])
            )
            fragments.extend(
                self._unique_str_list(edu.get("awards", []) if isinstance(edu.get("awards"), list) else [])
            )

        for work in parsed.get("work_experience", []):
            if not isinstance(work, dict):
                continue
            for field in ("company", "role", "start_date", "end_date"):
                value = self._safe_text(work.get(field))
                if value:
                    fragments.append(value)
            for resp in work.get("responsibilities", []):
                if not isinstance(resp, dict):
                    continue
                description = self._safe_text(resp.get("description"))
                if description:
                    fragments.append(description)
                fragments.extend(
                    self._unique_str_list(resp.get("metrics", []) if isinstance(resp.get("metrics"), list) else [])
                )
                fragments.extend(
                    self._unique_str_list(resp.get("keywords", []) if isinstance(resp.get("keywords"), list) else [])
                )

        for item in parsed.get("research", []):
            if not isinstance(item, dict):
                continue
            for field in ("title", "type", "start_date", "end_date", "description", "contributions"):
                value = self._safe_text(item.get(field))
                if value:
                    fragments.append(value)
            pubs = item.get("publications") if isinstance(item.get("publications"), list) else []
            for pub in pubs:
                if not isinstance(pub, dict):
                    continue
                for field in ("title", "venue", "status"):
                    value = self._safe_text(pub.get(field))
                    if value:
                        fragments.append(value)

        skills = parsed.get("skills", {})
        if isinstance(skills, dict):
            for value in skills.values():
                if not isinstance(value, list):
                    continue
                for entry in value:
                    if isinstance(entry, dict):
                        language = self._safe_text(entry.get("language"))
                        level = self._safe_text(entry.get("level"))
                        if language:
                            fragments.append(f"{language} {level}".strip())
                    else:
                        text = self._safe_text(entry)
                        if text:
                            fragments.append(text)

        return self._dedupe_lines(fragments)

    def _segment_line_for_match(self, line: str) -> list[str]:
        parts = re.split(
            r"[，。；：:,.!！?？·•\-/\\|()（）\[\]{}<>《》\"'`~@#$%^&*+=\s]+",
            line,
        )
        tokens: list[str] = []
        for part in parts:
            normalized = self._normalize_match_text(part)
            if len(normalized) >= 3:
                tokens.append(normalized)
        return tokens

    def _is_similar_to_fragments(
        self,
        line: str,
        normalized_line: str,
        normalized_fragments: list[str],
    ) -> bool:
        """Use fuzzy match to suppress false-positive unmapped lines."""
        if not normalized_line or not normalized_fragments:
            return False

        candidates = [normalized_line]
        for token in self._segment_line_for_match(line):
            if len(token) >= 6:
                candidates.append(token)

        deduped_candidates = self._dedupe_lines(candidates)
        for candidate in deduped_candidates:
            size = len(candidate)
            if size < 5:
                continue

            if size <= 8:
                threshold = 0.93
            elif size <= 16:
                threshold = 0.88
            else:
                threshold = 0.84

            for fragment in normalized_fragments:
                if not fragment:
                    continue
                if candidate in fragment or fragment in candidate:
                    return True

                size_gap = abs(len(fragment) - size)
                if size_gap > max(24, int(max(len(fragment), size) * 0.75)):
                    continue

                score = difflib.SequenceMatcher(None, candidate, fragment).ratio()
                if score >= threshold:
                    return True

        return False

    def _find_unmapped_lines(self, raw_text: str, parsed: dict) -> list[str]:
        """找出疑似未映射到结构化字段的原文行。"""
        fragments = self._collect_profile_fragments(parsed)
        normalized_fragments = self._dedupe_lines(
            [
                self._normalize_match_text(fragment)
                for fragment in fragments
                if len(self._normalize_match_text(fragment)) >= 3
            ]
        )
        haystack = "".join(normalized_fragments)
        if not haystack:
            return []

        fragment_tokens = set(normalized_fragments)

        raw_lines = self._dedupe_lines(raw_text.splitlines())
        unmatched: list[str] = []

        for line in raw_lines:
            normalized = self._normalize_match_text(line)
            if not normalized:
                continue
            # 过滤短标题行
            if len(normalized) <= 3 and not any(ch.isdigit() for ch in normalized):
                continue

            if normalized in haystack:
                continue

            tokens = self._segment_line_for_match(line)
            if tokens:
                hit_count = 0
                for token in tokens:
                    if token in haystack or token in fragment_tokens:
                        hit_count += 1
                # If most segments are already covered by parsed fields, treat as mapped.
                if hit_count / len(tokens) >= 0.6:
                    continue

            if self._is_similar_to_fragments(line, normalized, normalized_fragments):
                continue

            unmatched.append(line)

        return unmatched[:120]

    async def _structured_extraction(
        self,
        llm_service,
        raw_text: str,
        api_key: Optional[str] = None,
        provider: Optional[str] = None,
        model: Optional[str] = None,
        api_endpoint: Optional[str] = None,
    ) -> dict:
        """Convert resume raw text to structured profile JSON."""
        system_prompt = """You are a Resume Parser Agent. Convert resume text into structured JSON.
Rules:
1) Do not omit factual content.
2) Do not fabricate facts.
3) Preserve numbers, dates, names, companies, schools.
4) Put unresolved items in metadata.open_questions.
5) Output valid JSON only."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Parse the following resume text:\n\n{raw_text}"},
        ]

        try:
            first_result = await llm_service.complete_json(
                messages=messages,
                json_schema={},
                temperature=0.1,
                api_key=api_key,
                provider=provider,
                model=model,
                api_endpoint=api_endpoint,
            )
            parsed = self._ensure_profile_shape(first_result["parsed"])
        except Exception as exc:
            if not self._is_llm_json_parse_error(exc):
                raise
            logger.warning("Structured extraction JSON parse failed, using fallback parser: %s", exc)
            parsed = self._fallback_profile_from_text(raw_text, str(exc))

        # Second pass for completeness.
        try:
            audit_prompt = """You are a resume extraction auditor.
Given source text and first-pass JSON, return a completed JSON.
Only add missing facts from source text. Do not fabricate. Do not remove facts."""
            audit_user = f"Source text:\n{raw_text}\n\nFirst pass JSON:\n{parsed}"
            audit_result = await llm_service.complete_json(
                messages=[
                    {"role": "system", "content": audit_prompt},
                    {"role": "user", "content": audit_user},
                ],
                json_schema={},
                temperature=0.0,
                api_key=api_key,
                provider=provider,
                model=model,
                api_endpoint=api_endpoint,
            )
            audited = self._ensure_profile_shape(audit_result["parsed"])
            parsed = self._merge_profiles(parsed, audited)
        except Exception as exc:
            logger.warning("Profile audit pass skipped: %s", exc)

        self._redistribute_work_responsibilities_from_raw(parsed, raw_text)

        metadata = parsed.setdefault("metadata", {})
        if not isinstance(metadata, dict):
            metadata = {}
            parsed["metadata"] = metadata

        missing_fields = metadata.get("missing_fields")
        if not isinstance(missing_fields, list) or not missing_fields:
            metadata["missing_fields"] = self._compute_missing_fields(parsed)
        else:
            metadata["missing_fields"] = self._unique_str_list(missing_fields)

        open_questions = metadata.get("open_questions")
        if not isinstance(open_questions, list):
            open_questions = []
        else:
            open_questions = [
                text
                for text in self._unique_str_list(open_questions)
                if not re.search(
                    r"(llm reason:|fallback parsing was used|fallback mode|failed to parse llm json output|first output preview|review ai classified sections)",
                    text,
                    re.IGNORECASE,
                )
            ]

        unmapped_lines = self._find_unmapped_lines(raw_text, parsed)
        if unmapped_lines:
            metadata["unmapped_lines"] = unmapped_lines

        metadata["open_questions"] = self._unique_str_list(open_questions)

        confidence_value = self._recalculate_confidence(
            parsed=parsed,
            raw_text=raw_text,
            missing_fields=metadata.get("missing_fields", []),
            unmapped_lines=unmapped_lines,
        )
        metadata["extraction_confidence"] = confidence_value

        return {
            "profile": parsed,
            "confidence": metadata.get("extraction_confidence", 0.8),
            "raw_text": raw_text,
            "missing_fields": metadata.get("missing_fields", []),
            "open_questions": metadata.get("open_questions", []),
        }

    def _fallback_profile_from_text(self, raw_text: str, reason: str) -> dict:
        """Heuristic fallback when LLM output is not valid JSON."""
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        first_line = lines[0] if lines else ""

        email_match = re.search(
            r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
            raw_text,
        )
        phone_match = re.search(
            r"(\+?\d[\d\s\-()]{7,}\d)|((?:1[3-9]\d{9}))",
            raw_text,
        )
        email = email_match.group(0) if email_match else None
        phone = phone_match.group(0) if phone_match else None

        guessed_name = first_line if 1 <= len(first_line) <= 24 else "Name Pending"
        summary_source = " ".join(lines[:3]) if lines else raw_text.strip()
        summary = summary_source[:220] if summary_source else None

        return {
            "basics": {
                "name": guessed_name,
                "phone": phone,
                "email": email,
                "degree": None,
                "graduation_year": None,
                "target_role": None,
                "summary": summary,
            },
            "education": [],
            "work_experience": [],
            "research": [],
            "skills": {
                "data_analysis": [],
                "visualization": [],
                "programming": [],
                "ai_tools": [],
                "design": [],
                "media": [],
                "languages": [],
                "other": [],
            },
            "metadata": {
                "extraction_confidence": 0.25,
                "missing_fields": ["education", "work_experience", "skills"],
                "open_questions": [
                    "AI parsing used fallback mode. Please review imported sections manually.",
                ],
                "unmapped_lines": lines[:80],
            },
        }


# global singleton
parser_service = ParserService()
