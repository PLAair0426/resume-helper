"""导出服务 - 多格式简历导出"""
import logging
from pathlib import Path
from typing import Optional

from backend.core.config import settings

logger = logging.getLogger(__name__)


class ExportService:
    """简历导出：支持PDF/DOCX/MD/TXT"""

    # 已有的HTML模板映射
    TEMPLATE_FILES = {
        "A": "resume_optimized_v2.html",
        "B": "resume_template_B.html",
        "C": "resume_template_C.html",
        "D": "resume_template_D.html",
        "E": "resume_template_E.html",
    }

    async def export(self, version, format: str = "pdf", template_id: str = "A") -> Path:
        """
        导出简历文件。

        Args:
            version: ResumeVersion model instance
            format: pdf/docx/md/txt
            template_id: A/B/C/D/E（仅PDF格式使用）

        Returns:
            导出文件的Path
        """
        export_dir = settings.storage.export_dir / version.session_id
        export_dir.mkdir(parents=True, exist_ok=True)

        if format == "md":
            return await self._export_markdown(version, export_dir)
        elif format == "txt":
            return await self._export_txt(version, export_dir)
        elif format == "pdf":
            return await self._export_pdf(version, export_dir, template_id)
        elif format == "docx":
            return await self._export_docx(version, export_dir)
        else:
            raise ValueError(f"不支持的导出格式: {format}")

    async def _export_markdown(self, version, export_dir: Path) -> Path:
        """导出Markdown"""
        path = export_dir / f"resume_v{version.version_number}_{version.version_type}.md"
        path.write_text(version.content_md, encoding="utf-8")
        return path

    async def _export_txt(self, version, export_dir: Path) -> Path:
        """导出纯文本（去除Markdown标记）"""
        import re
        text = version.content_md
        # 简单去除Markdown标记
        text = re.sub(r'#{1,6}\s+', '', text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        text = re.sub(r'- ', '  ', text)

        path = export_dir / f"resume_v{version.version_number}_{version.version_type}.txt"
        path.write_text(text, encoding="utf-8")
        return path

    async def _export_pdf(self, version, export_dir: Path, template_id: str) -> Path:
        """导出PDF（通过HTML模板 + Playwright）"""
        from jinja2 import Template

        # 如果有预渲染的HTML，直接用
        if version.content_html:
            html_content = version.content_html
        else:
            # 使用Markdown转HTML的简单方案
            html_content = self._md_to_simple_html(version.content_md, template_id)

        # 写入临时HTML
        html_path = export_dir / f"temp_v{version.version_number}.html"
        html_path.write_text(html_content, encoding="utf-8")

        # Playwright生成PDF
        pdf_path = export_dir / f"resume_v{version.version_number}_{version.version_type}.pdf"
        await self._html_to_pdf(html_path, pdf_path)

        return pdf_path

    async def _html_to_pdf(self, html_path: Path, pdf_path: Path):
        """使用Playwright将HTML转为PDF"""
        from playwright.async_api import async_playwright

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page(viewport={"width": 794, "height": 1123})
            await page.goto(f"file:///{str(html_path).replace(chr(92), '/')}")
            await page.pdf(
                path=str(pdf_path),
                format="A4",
                print_background=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
            )
            await page.close()
            await browser.close()

    async def _export_docx(self, version, export_dir: Path) -> Path:
        """导出DOCX（ATS友好版）"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(10)

        # 解析Markdown并写入DOCX
        lines = version.content_md.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            else:
                # 去除加粗标记
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
                doc.add_paragraph(clean)

        path = export_dir / f"resume_v{version.version_number}_{version.version_type}.docx"
        doc.save(str(path))
        return path

    def _md_to_simple_html(self, md_content: str, template_id: str) -> str:
        """简单的Markdown转HTML（用于PDF导出）"""
        import re

        html_lines = []
        for line in md_content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                html_lines.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith("## "):
                html_lines.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("### "):
                html_lines.append(f"<h3>{line[4:]}</h3>")
            elif line.startswith("- "):
                html_lines.append(f"<p style='padding-left:8px;'>▸ {line[2:]}</p>")
            else:
                # 处理加粗
                line = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', line)
                html_lines.append(f"<p>{line}</p>")

        body = "\n".join(html_lines)
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>简历</title>
<style>
  @page {{ size: A4; margin: 0; }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: "Microsoft YaHei", sans-serif; color: #333; background: #fff;
         width: 210mm; height: 297mm; padding: 12mm 14mm; line-height: 1.5; font-size: 9px; overflow: hidden; }}
  h1 {{ font-size: 20px; text-align: center; margin-bottom: 8px; color: #1B2A4A; }}
  h2 {{ font-size: 11px; border-bottom: 1.5px solid #1B2A4A; padding-bottom: 2px; margin: 8px 0 4px; color: #1B2A4A; }}
  h3 {{ font-size: 9.5px; margin: 4px 0 2px; color: #333; }}
  p {{ margin-bottom: 2px; text-align: justify; }}
  strong {{ font-weight: 600; }}
</style></head>
<body>{body}</body></html>"""


# 全局单例
export_service = ExportService()
